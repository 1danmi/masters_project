from docx import Document

from bert2vec.data_process.tf_idf import getBOWSet
from bert2vec.data_process.bm_25 import get_BM25_Mat

# from bert2vec.model.dyvat.dyvat_bert2vec import Dyvat_bert2vec
# from shared_files.bert2vec.model.dyvat.dyvat_bert2vec import Dyvat_bert2vec

import numpy as np
import pandas as pd
from sklearn.metrics.pairwise import cosine_similarity

PATH = "_data_process"


# returns a matrix with the cosine similarity between each vector in mat
def getCosineMatrix(mat):
    N = len(mat)

    cos_mat = np.zeros((N, N))
    for i in range(N):
        for j in range(i):
            cos_mat[i][j] = cosine_similarity(
                np.reshape(pd.Series.to_numpy(mat.iloc[i]), (1, -1)),
                np.reshape(pd.Series.to_numpy(mat.iloc[j]), (1, -1)),
            )

    return cos_mat


def removeStopWords(rows):
    from nltk.corpus import stopwords

    stop_words = set(stopwords.words("english"))
    for r in rows:
        for sp in stop_words:
            if sp in r.BOW:
                r.BOW.pop(sp)
    return rows


def removePunctuation(rows):
    from string import punctuation

    for r in rows:
        for sp in punctuation:
            if sp in r.BOW:
                r.BOW.pop(sp)
    return rows


def getCosineSimMatrix(rows):
    df = pd.DataFrame([i.vec for i in rows])
    cosMat = getCosineMatrix(df)
    return cosMat


def getBM25_SimMatrix(rows, k1=1.5):
    bowSet = getBOWSet(rows)
    BM25_Cos = getCosineMatrix(get_BM25_Mat(bowSet, rows, k1))
    return BM25_Cos


# def getCosineMatrixes(rows, k1 = 1.5, b = 0):
#     df = pd.DataFrame([i.vec for i in rows])
#     cosMat = getCosineMatrix(df)
#     from copy import deepcopy
#     # rows_sw = removeStopWords(deepcopy(rows))
#     bowSet = getBOWSet(rows)
#     # TF_IDF_Cos = getCosineMatrix(get_TFIDF_Mat(bowSet, rows_sw))
#     # np.savetxt(f"tf_idf\\rows\\{rowsName}_cosineSim_TF_IDF.csv",TF_IDF_Cos,delimiter=' , ')
#     BM25_Cos = getCosineMatrix(get_BM25_Mat(bowSet, rows,k1, b))
#     # return cosMat, BM25_Cos
#     return cosMat, BM25_Cos


# receives a cosine similarity matrix and a BOW_matrix, returns a dictionary with
# {i : [j, cos_sim, bow_sim]} where cos_mat[i][j] >= cosThreshold and BOW_mat[i][j] >= bowThreshold
def getSimilarEntries(cosMat, bowMat, cosThreshold=0.85, bowThreshold=0.7) -> dict:
    vectors = dict()
    ANDMat = np.logical_and((cosMat >= cosThreshold), (bowMat >= bowThreshold))
    for i in range(len(ANDMat)):
        for j in range(i):
            # create a list of entries that are similar to the entry
            if ANDMat[i][j]:
                if j in vectors:
                    vectors[j].append((i, cosMat[i][j], bowMat[i][j]))
                else:
                    vectors[j] = [(i, cosMat[i][j], bowMat[i][j])]
    return vectors


# saves a report of all vectors
# vectors: {i :[j..]} where i is some entry number in rows
# and j are similar entries
# report of each entry: entry, count, sumOfBOW
# and the highest 120 BOW
def saveReport(rowsName, rows, vectors, k=1.5, b=0):
    doc = Document()
    doc.add_heading(rowsName, 0)
    # report of each row: entry, count, sumOfBOW
    # for each [row,column] == 1 : save the highest 120 BOW
    for i in vectors.keys():
        doc.add_heading(f"entry {i} similarities:", 1)
        doc.add_paragraph(
            f"entry {i}, count: {rows[i].count}, sumBOW: {sum(rows[i].BOW.values())}",
            style="List Bullet",
        )
        doc.add_paragraph("BOW:")
        doc.add_paragraph(str(sorted(rows[i].BOW.items(), key=lambda x: x[1], reverse=True)[:120]))
        for j, cosSim, bowSim in vectors[i]:
            doc.add_paragraph(
                f"entry {j}, count: {rows[j].count}, sumBOW: {sum(rows[j].BOW.values())}",
                style="List Bullet",
            )
            doc.add_paragraph(f"Cos_similarity: {cosSim}, bowSim: {bowSim}", style="List Bullet")
            doc.add_paragraph("BOW:")
            doc.add_paragraph(str(sorted(rows[j].BOW.items(), key=lambda x: x[1], reverse=True)[:120]))

    doc.save(f"{PATH}\\{rowsName}\\report_bm25_k{k}_b{b}_countDiv_no_stopWords.docx")


# saveDocReport('bank',bert2vec.rows['bank'])
def saveDocReport(rowsName, rows, k=1.5):
    import os

    if not os.path.exists(f"{PATH}\\{rowsName}"):
        os.makedirs(f"{PATH}\\{rowsName}")
    # get matrixes
    rows_sw = removeStopWords(rows)
    cosMat = getCosineSimMatrix(rows)
    BM25Mat = getBM25_SimMatrix(rows, k)
    # switch values >= .85 to 1 and < .85 to 0 in both
    np.savetxt(f"{PATH}\\{rowsName}\\vec_cosineSim.csv", cosMat, delimiter=" , ")
    np.savetxt(f"{PATH}\\{rowsName}\\BM_25_k{k}_cosineSim.csv", BM25Mat, delimiter=" , ")

    # preform AND operation between them

    vectors = getSimilarEntries(cosMat, BM25Mat)
    # save document
    saveReport(rowsName, rows_sw, vectors, k=k, b=0)


# def saveXlsDyvat(path : str, dyvat : Dyvat_bert2vec, words_BOW : tuple, numresults = 30):
#     import os
#     import string
#     dirpath = path + f"\\{words_BOW[0]}"
#     if not os.path.exists(dirpath): os.makedirs(dirpath)
#
#     #base algoritm:
#     base = dyvat.algorithm_base(numresults, (words_BOW[0], words_BOW[1]))
#     df = pd.DataFrame(columns=['word', 'bow'])
#     for i in range(len(base)):
#         df.loc[i] = [base[i].s, sorted([item for item in base[i].BOW.items() if item[0] not in string.punctuation],
#                                        key=lambda x: x[1], reverse=True)[:10]]
#     df.to_excel(dirpath + f"\\b2v_{words_BOW[0]}_base.xlsx")
#
#     #avg algoritm:
#     avg = dyvat.algorithm_avg(numresults, (words_BOW[0], words_BOW[1]))
#     df = pd.DataFrame(columns=['word', 'bow'])
#     for i in range(len(avg)):
#         df.loc[i] = [avg[i].s, sorted([item for item in avg[i].BOW.items() if item[0] not in string.punctuation],
#                                        key=lambda x: x[1], reverse=True)[:10]]
#     df.to_excel(dirpath + f"\\b2v_{words_BOW[0]}_avg.xlsx")
#     #
#     # #dyvat algoritm:
#     # dyv = dyvat.algorithm_dyvat(numresults, (words_BOW[0], words_BOW[1]), S=3)
#     # df = pd.DataFrame(columns=['word', 'bow'])
#     # for i in range(len(dyv)):
#     #     df.loc[i] = [dyv[i].s, sorted([item for item in dyv[i].BOW.items() if item[0] not in string.punctuation],
#     #                                    key=lambda x: x[1], reverse=True)[:10]]
#     # df.to_excel(dirpath + f"\\b2v_{words_BOW[0]}_dyvat_s{3}.xlsx")
#
# def saveDocDyvat(path : str, dyvat : Dyvat_bert2vec, words_BOW : tuple, numresults = 30):
#     import string
#     from docx import Document
#     doc = Document()
#     doc.add_heading(f"word: {words_BOW[0]}:")
#     doc.add_paragraph(f"base algorithm:" ,style='List Bullet')
#     base = dyvat.algorithm_base(numresults, (words_BOW[0], words_BOW[1]))
#     for i in range(len(base)):
#         doc.add_paragraph(f"{i}:\t{base[i].s}")
#         doc.add_paragraph(f"\t\tbow:{sorted([item for item in base[i].BOW.items() if item[0] not in string.punctuation],key=lambda x: x[1], reverse=True)[:10]}")
#
#     doc.add_paragraph(f"avg algorithm:" ,style='List Bullet')
#     avg = dyvat.algorithm_avg(numresults,(words_BOW[0], words_BOW[1]))
#     for i in range(len(avg)):
#         doc.add_paragraph(f"{i}:\t{avg[i].s}")
#         doc.add_paragraph(f"\t\tbow:{sorted([item for item in avg[i].BOW.items() if item[0] not in string.punctuation],key=lambda x: x[1], reverse=True)[:10]}")
#
#     doc.add_paragraph(f"dyvat algorithm:" ,style='List Bullet')
#     dyv = dyvat.algorithm_dyvat(numresults, (words_BOW[0], words_BOW[1]))
#     for i in range(len(dyv)):
#         doc.add_paragraph(f"{i}:\t{dyv[i].s}")
#         doc.add_paragraph(f"\t\tbow:{sorted([item for item in dyv[i].BOW.items() if item[0] not in string.punctuation],key=lambda x: x[1], reverse=True)[:10]}")
#
#     doc.save(f"{path}/dyvat_doc_{words_BOW[0]}.docx")
